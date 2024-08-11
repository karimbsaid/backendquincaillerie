from django.http import HttpResponse, Http404
from docx import Document
from CommandApi.models import Command  # Assuming Command is the model where your data is stored
import os
import io

def index(request):
    # Get the username from query parameters, if provided
    username_param = request.GET.get('username', None)
    print(username_param)
    
    # Retrieve commands from the database, filtered by username if provided
    if username_param:
        commands = Command.objects.filter(user__username=username_param).select_related('user', 'product')
        if not commands.exists():
            raise Http404(f"No commands found for user {username_param}")
    else:
        commands = Command.objects.all().select_related('user', 'product')
        if not commands.exists():
            raise Http404("No commands found")
    
    # Dictionary to store commands by user
    user_commands = {}
    
    for command in commands:
        username = command.user.username
        if username not in user_commands:
            user_commands[username] = []
        user_commands[username].append(command)
    
    # For simplicity, assuming one user at a time if username_param is provided
    if username_param:
        user_command_list = user_commands[username_param]
        document = Document()
        document.add_heading(f'Facture for {username_param}', 0)
        
        # Add a table with the user's commands
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Product ID'
        hdr_cells[2].text = 'Product Name'
        hdr_cells[3].text = 'Category'
        hdr_cells[4].text = 'Price'

        for command in user_command_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(command.quantity)
            row_cells[1].text = str(command.product.id)
            row_cells[2].text = command.product.name
            row_cells[3].text = command.product.category
            row_cells[4].text = str(command.product.price)
        
        # Save the document to a BytesIO object
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        # Create the HTTP response with the appropriate headers
        response = HttpResponse(doc_io.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=facture_{username_param}.docx'
        return response
    else:
        # Create a ZIP file with all factures if no specific user is mentioned
        import zipfile
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
            for username, user_command_list in user_commands.items():
                document = Document()
                document.add_heading(f'Facture for {username}', 0)
                
                # Add a table with the user's commands
                table = document.add_table(rows=1, cols=5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Qty'
                hdr_cells[1].text = 'Product ID'
                hdr_cells[2].text = 'Product Name'
                hdr_cells[3].text = 'Category'
                hdr_cells[4].text = 'Price'

                for command in user_command_list:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(command.quantity)
                    row_cells[1].text = str(command.product.id)
                    row_cells[2].text = command.product.name
                    row_cells[3].text = command.product.category
                    row_cells[4].text = str(command.product.price)
                
                # Save the document to a BytesIO object
                doc_io = io.BytesIO()
                document.save(doc_io)
                doc_io.seek(0)
                
                # Add the document to the ZIP file
                zip_file.writestr(f'facture_{username}.docx', doc_io.read())
        
        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer.read(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename=factures.zip'
        return response
